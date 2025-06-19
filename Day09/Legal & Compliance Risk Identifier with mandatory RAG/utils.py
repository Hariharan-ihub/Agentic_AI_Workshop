import os
import json
import numpy as np
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
from dataclasses import dataclass
import google.generativeai as genai
from langchain.text_splitter import RecursiveCharacterTextSplitter

@dataclass
class Document:
    page_content: str
    metadata: Dict[str, Any]

class VectorStore:
    def __init__(self, store_path: str = "vector_store"):
        self.store_path = store_path
        # Initialize Google Generative AI
        genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
        self.model = genai.GenerativeModel('gemini-pro')
        
        # Store documents and metadata
        self.documents = []
        self.vectors = []
        self.metadata = []
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )
        
        # Load existing data if available
        self._load_store()
    
    def _load_store(self):
        """Load existing vector store data"""
        if os.path.exists(f"{self.store_path}.json"):
            with open(f"{self.store_path}.json", 'r') as f:
                data = json.load(f)
                self.documents = data.get("documents", [])
                self.vectors = [np.array(v) for v in data.get("vectors", [])]
                self.metadata = data.get("metadata", [])
    
    def _save_store(self):
        """Save vector store data"""
        with open(f"{self.store_path}.json", 'w') as f:
            json.dump({
                "documents": self.documents,
                "vectors": [v.tolist() for v in self.vectors],
                "metadata": self.metadata
            }, f)
    
    def _get_embedding(self, text: str) -> np.ndarray:
        """Get embedding for a text using Google's embedding model"""
        response = genai.embed_content(
            model="models/embedding-001",
            content=text,
            task_type="retrieval_document",
            title="Document"
        )
        return np.array(response["embedding"])
    
    def _cosine_similarity(self, a: np.ndarray, b: np.ndarray) -> float:
        """Calculate cosine similarity between two vectors"""
        return np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b))
    
    def add_business_analysis(self, business_desc: str, analysis: Dict[str, Any]):
        """Add a business analysis to the store."""
        # Split text into chunks
        chunks = self.text_splitter.split_text(business_desc)
        
        # Add each chunk to store
        for chunk in chunks:
            vector = self._get_embedding(chunk)
            self.vectors.append(vector)
            self.documents.append(chunk)
            self.metadata.append({
                "type": "business_analysis",
                "analysis": analysis
            })
        
        # Save changes
        self._save_store()
    
    def add_risk_profile(self, domain: str, geography: str, profile: Dict[str, Any]):
        """Add a risk profile to the store."""
        key = f"{domain}_{geography}"
        vector = self._get_embedding(key)
        self.vectors.append(vector)
        self.documents.append(key)
        self.metadata.append({
            "type": "risk_profile",
            "domain": domain,
            "geography": geography,
            "profile": profile
        })
        
        # Save changes
        self._save_store()
    
    def add_legal_document(self, content: str, metadata: Dict[str, Any]) -> None:
        """Add a legal document to the vector store"""
        # Split text into chunks
        chunks = self.text_splitter.split_text(content)
        
        # Add each chunk to store
        for chunk in chunks:
            vector = self._get_embedding(chunk)
            self.vectors.append(vector)
            self.documents.append(chunk)
            self.metadata.append({
                "type": "legal_document",
                **metadata
            })
        
        # Save changes
        self._save_store()
    
    def similarity_search_with_score(self, query: str, k: int = 3) -> List[Tuple[Document, float]]:
        """Search for similar documents and return them with their similarity scores."""
        query_vector = self._get_embedding(query)
        
        # Calculate similarities
        similarities = [
            self._cosine_similarity(query_vector, vec)
            for vec in self.vectors
        ]
        
        # Get top k results
        top_k_indices = np.argsort(similarities)[-k:][::-1]
        
        results = []
        for idx in top_k_indices:
            results.append((
                Document(
                    page_content=self.documents[idx],
                    metadata=self.metadata[idx]
                ),
                float(similarities[idx])
            ))
        return results
    
    def search_legal_documents(self, query: str, jurisdiction: Optional[str] = None, k: int = 5) -> List[Dict[str, Any]]:
        """Search for legal documents in the vector store"""
        results = self.similarity_search_with_score(query, k=k)
        
        filtered_results = []
        for doc, score in results:
            if doc.metadata["type"] == "legal_document":
                if jurisdiction is None or doc.metadata.get("jurisdiction") == jurisdiction:
                    filtered_results.append({
                        "content": doc.page_content,
                        "metadata": doc.metadata,
                        "similarity": score
                    })
        
        return filtered_results
    
    def find_similar_business(self, business_desc: str, threshold: float = 0.8) -> Optional[Dict[str, Any]]:
        """Find similar business analysis."""
        results = self.similarity_search_with_score(business_desc, k=1)
        
        if results and results[0][1] >= threshold:
            return results[0][0].metadata["analysis"]
        return None
    
    def find_similar_risk_profile(self, domain: str, geography: str, threshold: float = 0.8) -> Optional[Dict[str, Any]]:
        """Find similar risk profile."""
        key = f"{domain}_{geography}"
        results = self.similarity_search_with_score(key, k=1)
        
        if results and results[0][1] >= threshold:
            return results[0][0].metadata["profile"]
        return None

def extract_text_from_document(file_content: bytes, file_type: str) -> str:
    """Extract text from uploaded document."""
    if file_type == "text/plain":
        return file_content.decode('utf-8')
    elif file_type == "application/pdf":
        from PyPDF2 import PdfReader
        from io import BytesIO
        pdf_file = BytesIO(file_content)
        pdf_reader = PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        from docx import Document
        from io import BytesIO
        doc_file = BytesIO(file_content)
        doc = Document(doc_file)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

def initialize_vector_store():
    """Initialize the vector store."""
    store = VectorStore()
    return store

def add_to_vector_store(vector_store, text, metadata):
    """
    Add a new document to the vector store with associated metadata.
    
    Args:
        vector_store: VectorStore instance
        text (str): The text to be embedded
        metadata (dict): Associated metadata for the text
    """
    vector_store.add_texts([text], [metadata])

def search_vector_store(vector_store: VectorStore, query: str, k: int = 3) -> List[Tuple[Document, float]]:
    """Search the vector store for similar documents."""
    return vector_store.similarity_search_with_score(query, k=k) 